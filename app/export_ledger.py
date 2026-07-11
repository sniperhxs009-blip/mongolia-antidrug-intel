"""Excel / Word / CSV 毒情台账导出。

修改原因：工程化——线下人工归档；支持 hide_details 分级脱敏。
"""
from __future__ import annotations

import csv
import zipfile
from pathlib import Path
from typing import Iterable, List
from xml.sax.saxutils import escape as xml_escape

from app.crawler.filters import sanitize_sensitive_text
from app.db.models import IntelItem


def _rows(items: Iterable[IntelItem], hide_details: bool) -> List[list]:
    out = []
    for it in items:
        title = sanitize_sensitive_text(it.title_zh or it.title or "", hide_details=hide_details)
        summary = sanitize_sensitive_text(it.summary_zh or it.summary or "", hide_details=hide_details)
        out.append(
            [
                it.id,
                title,
                it.org_name or "",
                it.category or "",
                it.intel_level or "",
                getattr(it, "credibility", "中") or "中",
                getattr(it, "alert_kind", "") or "",
                getattr(it, "port_tag", "") or "",
                getattr(it, "drug_type", "") or "",
                it.published_at.isoformat() if it.published_at else "",
                it.url or "",
                (summary or "")[:500],
            ]
        )
    return out


HEADERS = [
    "id", "title", "org", "category", "level", "credibility",
    "alert_kind", "port_tag", "drug_type", "published_at", "url", "summary",
]


def export_intel_csv(
    items: Iterable[IntelItem],
    path: str | Path,
    *,
    hide_details: bool = False,
) -> str:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(HEADERS)
        w.writerows(_rows(items, hide_details))
    return str(path)


def export_intel_xlsx(
    items: Iterable[IntelItem],
    path: str | Path,
    *,
    hide_details: bool = False,
) -> str:
    """无 openpyxl 依赖的简易 xlsx（SpreadsheetML zip）。"""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from openpyxl import Workbook  # type: ignore

        wb = Workbook()
        ws = wb.active
        ws.title = "毒情台账"
        ws.append(HEADERS)
        for row in _rows(items, hide_details):
            ws.append(row)
        wb.save(path)
        return str(path)
    except Exception:
        # 回退：写 CSV 同名改扩展，或最小 xlsx
        return _write_minimal_xlsx(path, HEADERS, _rows(items, hide_details))


def _write_minimal_xlsx(path: Path, headers: list, rows: list) -> str:
    def cell(v):
        return f'<c t="inlineStr"><is><t>{xml_escape(str(v))}</t></is></c>'

    sheet_rows = []
    for i, row in enumerate([headers] + rows, 1):
        cells = "".join(cell(c) for c in row)
        sheet_rows.append(f'<row r="{i}">{cells}</row>')
    sheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(sheet_rows)}</sheetData></worksheet>'
    )
    content_types = """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>"""
    wb = """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="毒情台账" sheetId="1" r:id="rId1"/></sheets>
</workbook>"""
    wb_rels = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>"""
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("xl/workbook.xml", wb)
        z.writestr("xl/_rels/workbook.xml.rels", wb_rels)
        z.writestr("xl/worksheets/sheet1.xml", sheet)
    return str(path)


def export_intel_docx(
    items: Iterable[IntelItem],
    path: str | Path,
    *,
    hide_details: bool = False,
) -> str:
    """简易 Word 文档（OOXML）；无 python-docx 时用最小 docx。"""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document  # type: ignore

        doc = Document()
        doc.add_heading("蒙古国禁毒情报台账", level=1)
        if hide_details:
            doc.add_paragraph("（分级导出：缴获数量与新型毒品细节已隐藏）")
        for row in _rows(items, hide_details):
            doc.add_paragraph(
                f"[{row[0]}] {row[1]} | {row[2]} | 可信度:{row[5]} | {row[9]} | {row[10]}"
            )
            if row[11]:
                doc.add_paragraph(row[11])
        doc.save(path)
        return str(path)
    except Exception:
        return _write_minimal_docx(path, _rows(items, hide_details), hide_details)


def _write_minimal_docx(path: Path, rows: list, hide_details: bool) -> str:
    paras = ['<w:p><w:r><w:t>蒙古国禁毒情报台账</w:t></w:r></w:p>']
    if hide_details:
        paras.append('<w:p><w:r><w:t>（分级导出：细节已隐藏）</w:t></w:r></w:p>')
    for row in rows:
        line = xml_escape(f"[{row[0]}] {row[1]} | {row[2]} | {row[10]}")
        paras.append(f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>")
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:body>{"".join(paras)}</w:body></w:document>'
    )
    content_types = """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document)
    return str(path)
